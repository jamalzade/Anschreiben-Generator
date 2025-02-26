import sys
import os
import re
import json
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from bs4 import BeautifulSoup
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QLabel, QLineEdit, QTextEdit, QPushButton, 
                            QMessageBox, QProgressBar, QFileDialog, QRadioButton,
                            QButtonGroup, QStackedWidget, QCalendarWidget, QDialog,
                            QDoubleSpinBox, QDateEdit, QComboBox, QInputDialog)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QDate
from PyQt6.QtGui import QFont, QTextCharFormat
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit

class WebScraper:
    @staticmethod
    def extract_job_description(url: str) -> str:
        try:
            # تنظیمات Chrome
            chrome_options = webdriver.ChromeOptions()
            chrome_options.add_argument('--headless')  # اجرا بدون نمایش مرورگر
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument("--disable-popup-blocking")
            chrome_options.add_argument("--disable-notifications")
            chrome_options.add_argument("--start-maximized")
            chrome_options.add_argument("--disable-infobars")
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            
            # راه‌اندازی Chrome
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)
            
            try:
                # تنظیم timeout برای لود شدن صفحه
                driver.set_page_load_timeout(30)
                
                # باز کردن صفحه
                driver.get(url)
                
                # صبر برای لود شدن محتوا (مثل کد اولیه)
                time.sleep(3)
                WebDriverWait(driver, 15).until(
                    lambda driver: driver.execute_script('return document.readyState') == 'complete'
                )
                
                job_desc = []
                
                # برای Indeed
                if 'indeed.com' in url or 'indeed.de' in url:
                    try:
                        job_desc_element = WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.ID, "jobDescriptionText"))
                        )
                        elements = job_desc_element.find_elements(By.CSS_SELECTOR, "p, li, div")
                        for element in elements:
                            text = element.text.strip()
                            if text and len(text) > 20:
                                job_desc.append(text)
                    except TimeoutException:
                        elements = driver.find_elements(By.CSS_SELECTOR, ".job-description p, .job-description li")
                        for element in elements:
                            text = element.text.strip()
                            if text and len(text) > 20:
                                job_desc.append(text)
                
                # برای LinkedIn
                elif 'linkedin.com' in url:
                    try:
                        job_desc_element = WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.CLASS_NAME, "show-more-less-html__markup"))
                        )
                        elements = job_desc_element.find_elements(By.CSS_SELECTOR, "p, li")
                        for element in elements:
                            text = element.text.strip()
                            if text:
                                job_desc.append(text)
                    except TimeoutException:
                        elements = driver.find_elements(By.CSS_SELECTOR, ".description__text p, .description__text li")
                        for element in elements:
                            text = element.text.strip()
                            if text:
                                job_desc.append(text)
                
                # برای Stepstone
                elif 'stepstone.de' in url:
                    try:
                        elements = driver.find_elements(By.CSS_SELECTOR, ".js-app-ld-ContentBlock p, .js-app-ld-ContentBlock li")
                        for element in elements:
                            text = element.text.strip()
                            if text:
                                job_desc.append(text)
                    except:
                        pass

                # برای سایر سایت‌ها
                if not job_desc:
                    elements = driver.find_elements(By.CSS_SELECTOR, "div p, div li")
                    for element in elements:
                        text = element.text.strip()
                        if text and len(text) > 20:
                            keywords = ['requirement', 'responsibility', 'qualification', 'experience',
                                      'skill', 'background', 'looking for', 'what you will do',
                                      'what you bring', 'your role', 'job description',
                                      'Anforderung', 'Qualifikation', 'Erfahrung', 'Aufgabe',
                                      'Verantwortung', 'Kenntnisse', 'Profil']
                            if any(keyword.lower() in text.lower() for keyword in keywords):
                                job_desc.append(text)
                
                if not job_desc:
                    print(f"No job description found. Elements checked: {len(elements)}")
                    raise ValueError(f"Could not find job description on this page. URL: {url}")
                
                return '\n'.join(f"• {item}" for item in job_desc)
            finally:
                driver.quit()
        except Exception as e:
            raise Exception(f"Error processing webpage: {str(e)}")

class FormValidator:
    @staticmethod
    def validate_email(email: str) -> bool:
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email))

    @staticmethod
    def validate_phone(phone: str) -> bool:
        pattern = r'^\+?[0-9\s-]{8,}$'
        return bool(re.match(pattern, phone))

class FormDataManager:
    def __init__(self, filename="form_data.json"):
        self.filename = filename

    def save_data(self, data: dict, profile_name: str):
        try:
            profiles = self.load_all_profiles()
            profiles[profile_name] = data
            with open(self.filename, 'w', encoding='utf-8') as f:
                json.dump(profiles, f, ensure_ascii=False, indent=4)
        except Exception as e:
            raise Exception(f"Failed to save form data: {str(e)}")

    def load_all_profiles(self) -> dict:
        try:
            if os.path.exists(self.filename):
                with open(self.filename, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return {}
        except Exception as e:
            raise Exception(f"Failed to load profiles: {str(e)}")

    def load_profile(self, profile_name: str) -> dict:
        profiles = self.load_all_profiles()
        return profiles.get(profile_name, {})

    def save_last_form_data(self, data: dict):
        """ذخیره آخرین داده‌های فرم در فایل JSON"""
        try:
            with open(self.filename.replace('.json', '_last.json'), 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
        except Exception as e:
            raise Exception(f"Failed to save last form data: {str(e)}")

    def load_last_form_data(self) -> dict:
        """لود آخرین داده‌های فرم از فایل JSON"""
        try:
            filename = self.filename.replace('.json', '_last.json')
            if os.path.exists(filename):
                with open(filename, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return {}
        except Exception as e:
            raise Exception(f"Failed to load last form data: {str(e)}")

class FileReader:
    @staticmethod
    def read_file(file_path: str) -> str:
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_extension in ['.doc', '.docx']:
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise ValueError("Unsupported file format")

class AnschreibenGenerator:
    def __init__(self):
        self.template_version = 1
    
    def toggle_template(self):
        self.template_version = 2 if self.template_version == 1 else 1
        
    def generate_anschreiben(self, full_name: str, email: str, phone: str,
                            address: str, city: str, company_name: str,
                            contact_person: str, company_address: str,
                            company_city: str, position: str,
                            responsibilities: str, start_date: str) -> str:
        if self.template_version == 1:
            return self._generate_template_one(full_name, email, phone, address, city,
                                            company_name, contact_person, company_address,
                                            company_city, position, responsibilities,
                                            start_date)
        else:
            return self._generate_template_two(full_name, email, phone, address, city,
                                            company_name, contact_person, company_address,
                                            company_city, position, responsibilities,
                                            start_date)

    def _generate_template_one(self, full_name: str, email: str, phone: str,
                            address: str, city: str, company_name: str,
                            contact_person: str, company_address: str,
                            company_city: str, position: str,
                            responsibilities: str, start_date: str) -> str:
        current_date = datetime.now().strftime("%d.%m.%Y")
        
        # استفاده از یک متن عمومی‌تر و خلاصه‌تر
        general_skills = """
• Mehrjährige Erfahrung in relevanten Bereichen
• Starke analytische und kommunikative Fähigkeiten
• Teamfähigkeit und Eigeninitiative
• Schnelle Einarbeitung in neue Aufgaben"""
        
        letter = f"""
{full_name}
{address}
{city}

{company_name}
{contact_person}
{company_address}
{company_city}

{current_date}

Bewerbung als {position}

Sehr geehrte(r) {contact_person},

mit großem Interesse habe ich Ihre Stellenausschreibung für die Position als {position} bei {company_name} gelesen. Mit meiner Erfahrung und meinen Fähigkeiten bin ich überzeugt, einen wertvollen Beitrag zu Ihrem Team leisten zu können.

Ich bringe folgende Qualifikationen mit:
{general_skills}

{company_name} beeindruckt mich durch seinen Ruf als innovatives Unternehmen. Ich freue mich darauf, meine Fähigkeiten in Ihrem Team einzubringen und zur Weiterentwicklung beizutragen.

Ich bin ab dem {start_date} oder nach Vereinbarung verfügbar.

Gerne überzeuge ich Sie in einem persönlichen Gespräch von meinen Qualifikationen. Ich freue mich auf Ihre Rückmeldung.

Mit freundlichen Grüßen
{full_name}

Kontakt:
E-Mail: {email}
Tel.: {phone}

Anlagen:
• Lebenslauf
• Zeugnisse
• Referenzen"""
        return letter

    def _generate_template_two(self, full_name: str, email: str, phone: str,
                            address: str, city: str, company_name: str,
                            contact_person: str, company_address: str,
                            company_city: str, position: str,
                            responsibilities: str, start_date: str) -> str:
        current_date = datetime.now().strftime("%d.%m.%Y")
        
        # استفاده از یک متن عمومی‌تر و خلاصه‌تر
        general_skills = """
• Erfahrung in relevanten Projekten
• Gute Kommunikations- und Teamfähigkeiten
• Hohe Lernbereitschaft und Flexibilität"""
        
        letter = f"""
Bewerbung: {position} - {company_name}

{full_name}
{address}
{city}
Tel.: {phone}
E-Mail: {email}

An
{company_name}
z.Hd. {contact_person}
{company_address}
{company_city}

{current_date}

Bewerbung für die Position als {position}

Sehr geehrte(r) {contact_person},

Ihre Stellenausschreibung für die Position als {position} bei {company_name} hat mein Interesse geweckt. Ich bin überzeugt, dass meine Erfahrungen und Fähigkeiten gut zu Ihren Anforderungen passen.

Meine Kernkompetenzen umfassen:
{general_skills}

Ich schätze den Ruf von {company_name} als führendes Unternehmen und möchte Teil Ihres Teams werden. Ich bin ab {start_date} oder nach Vereinbarung verfügbar.

Ich freue mich auf die Gelegenheit, Sie in einem Gespräch von meinen Qualifikationen zu überzeugen.

Mit freundlichen Grüßen

{full_name}

Anlagen:
- Lebenslauf
- Zeugnisse
- Referenzen"""
        return letter

    def save_as_word(self, letter: str, filename: str):
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # تغییر فونت به Calibri برای حرفه‌ای‌تر شدن
        font.size = Pt(12)  # افزایش اندازه به 12pt برای خوانایی بهتر
        
        # تنظیم فاصله‌های پاراگراف‌ها
        paragraph_spacing = doc.styles['Normal'].paragraph_format
        paragraph_spacing.space_after = Pt(6)  # فاصله بعد از هر پاراگراف
        
        paragraphs = letter.split('\n')
        for p in paragraphs:
            if p.strip():
                doc_p = doc.add_paragraph(p)
                if p.startswith('•'):
                    doc_p.paragraph_format.left_indent = Pt(24)  # تورنتگی بیشتر برای لیست‌ها
                elif p.startswith(('Sehr geehrte', 'Bewerbung')):
                    doc_p.style = doc.styles['Heading 2']  # تیترها با استایل بولد
                    doc_p.paragraph_format.space_after = Pt(12)  # فاصله بیشتر بعد از تیترها
                else:
                    doc_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # چیدمان چپ
        
        # تنظیم حاشیه‌ها
        doc.sections[0].page_margins.left = Pt(70)  # 2.5cm
        doc.sections[0].page_margins.right = Pt(70)  # 2.5cm
        doc.sections[0].page_margins.top = Pt(70)  # 2.5cm
        doc.sections[0].page_margins.bottom = Pt(70)  # 2.5cm
        
        doc.save(filename)

class GenerateLetterThread(QThread):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    
    def __init__(self, generator, data):
        super().__init__()
        self.generator = generator
        self.data = data
    
    def run(self):
        try:
            result = self.generator.generate_anschreiben(**self.data)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.generator = AnschreibenGenerator()
        self.data_manager = FormDataManager()
        self.init_ui()
        # لود خودکار آخرین داده‌های فرم
        self.load_last_form_data()

    def _add_field(self, layout, label):
        field_layout = QHBoxLayout()
        field_layout.addWidget(QLabel(label))
        line_edit = QLineEdit()
        field_layout.addWidget(line_edit)
        layout.addLayout(field_layout)
        return line_edit

    def save_form_data(self):
        try:
            profile_name = self.profile_combo.currentText()
            if profile_name == "New Profile":
                profile_name, ok = QInputDialog.getText(self, 'Profile Name', 'Enter a name for this profile:')
                if not ok or not profile_name:
                    return
                self.profile_combo.addItem(profile_name)
                self.profile_combo.setCurrentText(profile_name)
            
            data = {
                'full_name': self.full_name.text(),
                'email': self.email.text(),
                'phone': self.phone.text(),
                'address': self.address.text(),
                'city': self.city.text(),
                'company_name': self.company_name.text(),
                'contact_person': self.contact_person.text(),
                'company_address': self.company_address.text(),
                'company_city': self.company_city.text(),
                'position': self.position.text(),
                'start_date': self.start_date.date().toString(Qt.DateFormat.ISODate),
                'responsibilities': self.responsibilities.toPlainText()
            }
            self.data_manager.save_data(data, profile_name)
            # ذخیره خودکار آخرین داده‌های فرم
            self.data_manager.save_last_form_data(data)
            QMessageBox.information(self, 'Success', f'Profile "{profile_name}" saved!')
        except Exception as e:
            QMessageBox.critical(self, 'Error', f'Failed to save profile: {str(e)}')

    def load_saved_profiles(self):
        try:
            profiles = self.data_manager.load_all_profiles()
            for profile_name in profiles.keys():
                self.profile_combo.addItem(profile_name)
        except Exception as e:
            QMessageBox.warning(self, 'Warning', f'Failed to load profiles: {str(e)}')

    def load_selected_profile(self, profile_name):
        if profile_name != "New Profile":
            try:
                data = self.data_manager.load_profile(profile_name)
                if data:
                    self.full_name.setText(data.get('full_name', ''))
                    self.email.setText(data.get('email', ''))
                    self.phone.setText(data.get('phone', ''))
                    self.address.setText(data.get('address', ''))
                    self.city.setText(data.get('city', ''))
                    self.company_name.setText(data.get('company_name', ''))
                    self.contact_person.setText(data.get('contact_person', ''))
                    self.company_address.setText(data.get('company_address', ''))
                    self.company_city.setText(data.get('company_city', ''))
                    self.position.setText(data.get('position', ''))
                    if 'start_date' in data:
                        self.start_date.setDate(QDate.fromString(data['start_date'], Qt.DateFormat.ISODate))
                    self.responsibilities.setPlainText(data.get('responsibilities', ''))
            except Exception as e:
                QMessageBox.warning(self, 'Warning', f'Failed to load profile: {str(e)}')
        else:
            self.reset_form()

    def load_last_form_data(self):
        """لود آخرین داده‌های فرم هنگام باز شدن برنامه"""
        try:
            last_data = self.data_manager.load_last_form_data()
            if last_data:
                self.full_name.setText(last_data.get('full_name', ''))
                self.email.setText(last_data.get('email', ''))
                self.phone.setText(last_data.get('phone', ''))
                self.address.setText(last_data.get('address', ''))
                self.city.setText(last_data.get('city', ''))
                self.company_name.setText(last_data.get('company_name', ''))
                self.contact_person.setText(last_data.get('contact_person', ''))
                self.company_address.setText(last_data.get('company_address', ''))
                self.company_city.setText(last_data.get('company_city', ''))
                self.position.setText(last_data.get('position', ''))
                if 'start_date' in last_data:
                    self.start_date.setDate(QDate.fromString(last_data['start_date'], Qt.DateFormat.ISODate))
                self.responsibilities.setPlainText(last_data.get('responsibilities', ''))
        except Exception as e:
            QMessageBox.warning(self, 'Warning', f'Failed to load last form data: {str(e)}')

    def fetch_from_url(self):
        url = self.url_input.text().strip()
        if not url:
            QMessageBox.warning(self, 'Input Error', 'Please enter a URL.')
            return
        if not any(domain in url.lower() for domain in ['indeed.com', 'indeed.de', 'linkedin.com', 'stepstone.de']):
            QMessageBox.warning(self, 'Input Error', 'Please enter a valid job posting URL (e.g., Indeed, LinkedIn, Stepstone).')
            return
        
        try:
            self.progress.setVisible(True)
            self.progress.setRange(0, 0)
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Icon.Information)
            msg.setText("Fetching job description...\nThis might take a few seconds.")
            msg.setStandardButtons(QMessageBox.StandardButton.NoButton)
            msg.show()
            description = WebScraper.extract_job_description(url)
            msg.close()
            preview = QMessageBox()
            preview.setIcon(QMessageBox.Icon.Information)
            preview.setText("Job description fetched successfully. Would you like to preview it?")
            preview.setDetailedText(description)
            preview.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if preview.exec() == QMessageBox.StandardButton.Yes:
                self.responsibilities.setPlainText(description)
        except Exception as e:
            QMessageBox.critical(self, 'Error', f'Failed to fetch job description: {str(e)}')
        finally:
            self.progress.setVisible(False)

    def init_ui(self):
        self.setWindowTitle('Anschreiben Generator')
        self.setMinimumWidth(900)
        
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        
        profile_layout = QHBoxLayout()
        profile_layout.addWidget(QLabel('Select Profile:'))
        self.profile_combo = QComboBox()
        self.profile_combo.addItem("New Profile")
        self.profile_combo.currentTextChanged.connect(self.load_selected_profile)
        profile_layout.addWidget(self.profile_combo)
        layout.addLayout(profile_layout)

        personal_layout = QVBoxLayout()
        company_layout = QVBoxLayout()
        job_layout = QVBoxLayout()
        
        personal_layout.addWidget(QLabel('Personal Information:'))
        self.full_name = self._add_field(personal_layout, 'Full Name:')
        self.email = self._add_field(personal_layout, 'Email:')
        self.phone = self._add_field(personal_layout, 'Phone:')
        self.address = self._add_field(personal_layout, 'Street Address:')
        self.city = self._add_field(personal_layout, 'City:')
        
        company_layout.addWidget(QLabel('Company Information:'))
        self.company_name = self._add_field(company_layout, 'Company Name:')
        self.contact_person = self._add_field(company_layout, 'Contact Person:')
        self.company_address = self._add_field(company_layout, 'Company Address:')
        self.company_city = self._add_field(company_layout, 'Company City:')
        
        job_layout.addWidget(QLabel('Job Information:'))
        self.position = self._add_field(job_layout, 'Position:')
        
        date_layout = QHBoxLayout()
        date_layout.addWidget(QLabel('Start Date:'))
        self.start_date = QDateEdit()
        self.start_date.setCalendarPopup(True)
        self.start_date.setDate(QDate.currentDate())
        date_layout.addWidget(self.start_date)
        job_layout.addLayout(date_layout)
        
        resp_method_layout = QHBoxLayout()
        resp_method_layout.addWidget(QLabel('Enter responsibilities manually (optional):'))
        self.responsibilities = QTextEdit()
        self.responsibilities.setMinimumHeight(100)
        self.responsibilities.setPlaceholderText("Optional: Enter general skills or responsibilities (e.g., 'Erfahrung in Projektmanagement, Teamfähigkeit')")
        job_layout.addWidget(self.responsibilities)
        
        # بخش اختیاری URL برای اطلاعات اضافی
        url_widget = QWidget()
        url_layout = QVBoxLayout(url_widget)
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText('Optional: Enter job posting URL for reference...')
        fetch_btn = QPushButton('Fetch Job Description (Optional)')
        fetch_btn.clicked.connect(self.fetch_from_url)
        url_layout.addWidget(QLabel('Job Posting URL (Optional):'))
        url_layout.addWidget(self.url_input)
        url_layout.addWidget(fetch_btn)
        job_layout.addWidget(url_widget)
        
        layout.addLayout(personal_layout)
        layout.addLayout(company_layout)
        layout.addLayout(job_layout)
        
        buttons_layout = QHBoxLayout()
        self.progress = QProgressBar()
        self.progress.setVisible(False)
        layout.addWidget(self.progress)
        
        self.generate_btn = QPushButton('Generate Anschreiben')
        self.generate_btn.clicked.connect(self.generate_letter)
        buttons_layout.addWidget(self.generate_btn)
        
        self.new_template_btn = QPushButton('Generate Alternative Template')
        self.new_template_btn.clicked.connect(self.generate_new_template)
        buttons_layout.addWidget(self.new_template_btn)
        
        self.preview_btn = QPushButton('Preview Letter')
        self.preview_btn.clicked.connect(self.preview_letter)
        self.preview_btn.setEnabled(False)
        buttons_layout.addWidget(self.preview_btn)
        
        self.reset_btn = QPushButton('Reset Form')
        self.reset_btn.clicked.connect(self.reset_form)
        buttons_layout.addWidget(self.reset_btn)
        
        layout.addLayout(buttons_layout)
        
        save_buttons_layout = QHBoxLayout()
        self.save_txt_btn = QPushButton('Save as Text')
        self.save_txt_btn.clicked.connect(lambda: self.save_letter('txt'))
        self.save_txt_btn.setEnabled(False)
        save_buttons_layout.addWidget(self.save_txt_btn)
        
        self.save_docx_btn = QPushButton('Save as Word')
        self.save_docx_btn.clicked.connect(lambda: self.save_letter('docx'))
        self.save_docx_btn.setEnabled(False)
        save_buttons_layout.addWidget(self.save_docx_btn)

        self.save_pdf_btn = QPushButton('Save as PDF')
        self.save_pdf_btn.clicked.connect(lambda: self.save_letter('pdf'))
        self.save_pdf_btn.setEnabled(False)
        save_buttons_layout.addWidget(self.save_pdf_btn)
        
        layout.addLayout(save_buttons_layout)
        
        self.save_form_btn = QPushButton('Save Form Data')
        self.save_form_btn.clicked.connect(self.save_form_data)
        layout.addWidget(self.save_form_btn)

    def closeEvent(self, event):
        """ذخیره خودکار داده‌ها هنگام بستن پنجره"""
        try:
            data = {
                'full_name': self.full_name.text(),
                'email': self.email.text(),
                'phone': self.phone.text(),
                'address': self.address.text(),
                'city': self.city.text(),
                'company_name': self.company_name.text(),
                'contact_person': self.contact_person.text(),
                'company_address': self.company_address.text(),
                'company_city': self.company_city.text(),
                'position': self.position.text(),
                'start_date': self.start_date.date().toString(Qt.DateFormat.ISODate),
                'responsibilities': self.responsibilities.toPlainText()
            }
            self.data_manager.save_last_form_data(data)
            event.accept()
        except Exception as e:
            QMessageBox.critical(self, 'Error', f'Failed to save form data on close: {str(e)}')
            event.accept()

    def reset_form(self):
        self.full_name.clear()
        self.email.clear()
        self.phone.clear()
        self.address.clear()
        self.city.clear()
        self.company_name.clear()
        self.contact_person.clear()
        self.company_address.clear()
        self.company_city.clear()
        self.position.clear()
        self.responsibilities.clear()
        self.url_input.clear()
        self.start_date.setDate(QDate.currentDate())
        self.save_txt_btn.setEnabled(False)
        self.save_docx_btn.setEnabled(False)
        self.save_pdf_btn.setEnabled(False)
        self.preview_btn.setEnabled(False)
        self.progress.setVisible(False)
        self.generate_btn.setEnabled(True)
        QMessageBox.information(self, 'Reset', 'All fields have been cleared.')
    
    def get_responsibilities(self) -> str:
        return self.responsibilities.toPlainText() or ""

    def generate_letter(self):
        try:
            responsibilities = self.get_responsibilities()
        except Exception as e:
            QMessageBox.warning(self, 'Input Error', str(e))
            return
            
        if not FormValidator.validate_email(self.email.text()):
            QMessageBox.warning(self, 'Input Error', 'Please enter a valid email address.')
            return
        if not FormValidator.validate_phone(self.phone.text()):
            QMessageBox.warning(self, 'Input Error', 'Please enter a valid phone number.')
            return
            
        data = {
            'full_name': self.full_name.text(),
            'email': self.email.text(),
            'phone': self.phone.text(),
            'address': self.address.text(),
            'city': self.city.text(),
            'company_name': self.company_name.text(),
            'contact_person': self.contact_person.text(),
            'company_address': self.company_address.text(),
            'company_city': self.company_city.text(),
            'position': self.position.text(),
            'responsibilities': responsibilities,
            'start_date': self.start_date.date().toString("dd.MM.yyyy")
        }
        
        if not all(value for key, value in data.items()):
            QMessageBox.warning(self, 'Input Error', 'Please fill in all fields.')
            return
            
        self.progress.setVisible(True)
        self.progress.setRange(0, 0)
        self.generate_btn.setEnabled(False)
        
        self.thread = GenerateLetterThread(self.generator, data)
        self.thread.finished.connect(self.handle_generation_complete)
        self.thread.error.connect(self.handle_generation_error)
        self.thread.start()

    def handle_generation_complete(self, result):
        self.progress.setVisible(False)
        self.generate_btn.setEnabled(True)
        self.save_txt_btn.setEnabled(True)
        self.save_docx_btn.setEnabled(True)
        self.save_pdf_btn.setEnabled(True)
        self.preview_btn.setEnabled(True)
        self.current_letter = result
        QMessageBox.information(self, 'Success', 'Letter generated successfully!\nYou can now preview or save it.')

    def handle_generation_error(self, error):
        QMessageBox.critical(self, 'Error', f'Failed to generate letter: {error}')
        self.progress.setVisible(False)
        self.generate_btn.setEnabled(True)

    def generate_new_template(self):
        self.generator.toggle_template()
        self.generate_letter()
        self.generator.toggle_template()

    def preview_letter(self):
        if hasattr(self, 'current_letter') and self.current_letter:
            preview_window = QDialog(self)
            preview_window.setWindowTitle('Letter Preview')
            preview_window.setMinimumSize(600, 400)
            
            layout = QVBoxLayout(preview_window)
            text_edit = QTextEdit()
            text_edit.setReadOnly(True)
            text_edit.setFont(QFont("Calibri", 12))  # تنظیم فونت و اندازه برای پیش‌نمایش
            text_edit.setPlainText(self.current_letter)
            layout.addWidget(text_edit)
            
            close_btn = QPushButton('Close')
            close_btn.clicked.connect(preview_window.accept)
            layout.addWidget(close_btn)
            
            preview_window.exec()
        else:
            QMessageBox.warning(self, 'No Letter', 'Please generate a letter first.')

    def save_letter(self, format_type='txt'):
        if not hasattr(self, 'current_letter'):
            QMessageBox.warning(self, 'No Letter', 'Please generate a letter first.')
            return
        
        if format_type == 'txt':
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "Save Anschreiben as Text",
                f"anschreiben_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                "Text Files (*.txt);;All Files (*)"
            )
            if file_name:
                try:
                    # اضافه کردن فاصله‌های منظم
                    formatted_letter = self.current_letter.replace('\n', '\n\n').replace('• ', '    • ')
                    with open(file_name, 'w', encoding='utf-8') as f:
                        f.write(formatted_letter)
                    QMessageBox.information(self, 'Success', f'Letter saved to {file_name}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Failed to save file: {str(e)}')
        
        elif format_type == 'docx':
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "Save Anschreiben as Word Document",
                f"anschreiben_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx);;All Files (*)"
            )
            if file_name:
                try:
                    self.generator.save_as_word(self.current_letter, file_name)
                    QMessageBox.information(self, 'Success', f'Letter saved to {file_name}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Failed to save Word document: {str(e)}')

        elif format_type == 'pdf':
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "Save Anschreiben as PDF",
                f"anschreiben_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "PDF Files (*.pdf);;All Files (*)"
            )
            if file_name:
                try:
                    temp_html = file_name + '.temp.html'
                    with open(temp_html, 'w', encoding='utf-8') as f:
                        # پردازش متن برای اضافه کردن تیترها
                        paragraphs = self.current_letter.split('\n')
                        html_content = "<html><head><meta charset='UTF-8'><style>body { font-family: Calibri, sans-serif; font-size: 12pt; line-height: 1.5; max-width: 800px; margin: auto; padding: 20px; }.bullet-point { margin-left: 24px; margin-top: 4px; margin-bottom: 4px; } h1, h2 { font-weight: bold; margin-bottom: 12px; }</style><body>"
                        for p in paragraphs:
                            if p.strip():
                                if p.startswith(('Sehr geehrte', 'Bewerbung')):
                                    html_content += f"<h2>{p}</h2>"
                                elif p.startswith(('•')):
                                    html_content += f"<div class='bullet-point'>{p}</div>"
                                else:
                                    html_content += f"<p>{p}</p>"
                        html_content += "</body></html>"
                        f.write(html_content)
                    options = {
                        'encoding': 'UTF-8',
                        'page-size': 'A4',
                        'margin-top': '25mm',  # افزایش به 2.5cm برای استاندارد آلمانی
                        'margin-right': '25mm',
                        'margin-bottom': '25mm',
                        'margin-left': '25mm'
                    }
                    pdfkit.from_file(temp_html, file_name, options=options)
                    os.remove(temp_html)
                    QMessageBox.information(self, 'Success', f'Letter saved to {file_name}')
                except Exception as e:
                    QMessageBox.critical(self, 'Error', f'Failed to save PDF file: {str(e)}')

def main():
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()